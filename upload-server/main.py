"""
Hermes WebUI 的 PDF / docx 上傳處理 server。

部署位置：mac mini（/Users/young/.hermes-upload-server/）
跑在：localhost:9120
被誰呼叫：ruru-hermes-ai-agent 前端，透過 SSH tunnel + Vite proxy（/upload-api）

為什麼需要這個 server？
  ChatPage 的純文字檔可以瀏覽器 file.text() 自己讀完，但 PDF 跟 docx
  需要重的解析（pymupdf / python-docx）。要前端做 = 拉 pdfjs-dist 進來，
  品質遠遜於 pymupdf 而且打包體積大。要呼叫 Hermes 後端的 ocr-and-documents
  skill = 動 NousResearch upstream 的程式碼，沒辦法持續維護。
  所以在這個 repo 自己起一個輕量 mini server，跟 Hermes 解耦。

路由：
  GET  /healthz                 — 不需 token，給 tunnel / 部署檢查用
  POST /upload                  — 上傳檔案，回傳抽出來的文字
  GET  /list                    — 列出 ruru 論文資料夾裡所有檔
  GET  /file/{filename}/text    — 重讀已上傳檔案的文字（不重新解析）

認證：
  所有 POST / 列檔 / 讀檔都要 Authorization: Bearer <token>
  token 在啟動時從 ~/.hermes-upload-server.token 讀（沒有就自動生成）
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
import secrets
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Annotated

import fitz  # PyMuPDF — 套件名 pymupdf，import 名是 fitz
from docx import Document
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# ─── 設定 ───────────────────────────────────────────────────────────

UPLOAD_ROOT = Path("/Users/young/Desktop/ruru_論文專區")
DEFAULT_CATEGORY = "uploads"
TOKEN_PATH = Path.home() / ".hermes-upload-server.token"
MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
ALLOWED_ORIGINS = [
    "http://localhost:5174",  # Vite dev server
    "http://127.0.0.1:5174",
]

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("upload-server")


# ─── Token 管理 ─────────────────────────────────────────────────────


def load_or_create_token() -> str:
    if TOKEN_PATH.exists():
        token = TOKEN_PATH.read_text().strip()
        if token:
            return token
    token = secrets.token_urlsafe(32)
    TOKEN_PATH.write_text(token)
    TOKEN_PATH.chmod(0o600)
    log.info("Generated new auth token at %s", TOKEN_PATH)
    return token


AUTH_TOKEN = load_or_create_token()


def require_auth(authorization: Annotated[str | None, Header()] = None) -> None:
    """Bearer token auth。tunnel 已經把 server 限制在 127.0.0.1 可達，
    這層 token 是縱深防禦：萬一 tunnel 設定錯，也擋掉外部請求。"""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Missing Bearer token",
        )
    # 使用 secrets.compare_digest 避免 timing attack
    if not secrets.compare_digest(authorization.removeprefix("Bearer "), AUTH_TOKEN):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid token",
        )


# ─── 檔案處理 ───────────────────────────────────────────────────────


SAFE_FILENAME_RE = re.compile(r"[^\w\-. 一-鿿]+")


def sanitize_filename(name: str) -> str:
    """擋 path traversal / null bytes / 控制字元；保留 CJK 字。"""
    name = unicodedata.normalize("NFC", name)
    name = name.replace("\x00", "").replace("/", "_").replace("\\", "_")
    # 去掉前置的 . 避免變成 hidden file
    name = name.lstrip(".")
    name = SAFE_FILENAME_RE.sub("_", name)
    name = name.strip() or "untitled"
    # 限制長度：macOS HFS+ 單個 path component 上限 255 bytes
    if len(name.encode("utf-8")) > 200:
        stem, _, ext = name.rpartition(".")
        ext = ("." + ext) if "." in name else ""
        name = stem[:100] + ext
    return name


def reserve_path(directory: Path, filename: str) -> Path:
    """同檔名加 (1) (2) ... 後綴，避免覆蓋既有檔。"""
    candidate = directory / filename
    if not candidate.exists():
        return candidate
    stem = candidate.stem
    suffix = candidate.suffix
    i = 1
    while True:
        candidate = directory / f"{stem} ({i}){suffix}"
        if not candidate.exists():
            return candidate
        i += 1


def extract_pdf(data: bytes) -> str:
    """pymupdf 抽 PDF 純文字。一頁一頁拼，頁與頁之間加分頁標記方便 AI 知道斷點。"""
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF 解析失敗：{e}") from e
    parts: list[str] = []
    for i, page in enumerate(doc, start=1):
        text = page.get_text("text") or ""
        parts.append(f"--- 第 {i} 頁 ---\n{text.strip()}")
    doc.close()
    return "\n\n".join(parts).strip()


def extract_docx(data: bytes) -> str:
    """python-docx 抽 docx 文字。paragraphs + tables 都收。"""
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"docx 解析失敗：{e}") from e
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract(data: bytes, ext: str) -> str:
    if ext == ".pdf":
        return extract_pdf(data)
    if ext in (".docx", ".doc"):
        return extract_docx(data)
    raise HTTPException(status_code=400, detail=f"不支援的副檔名：{ext}")


# ─── App ────────────────────────────────────────────────────────────


app = FastAPI(title="Hermes WebUI Upload Server", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["GET", "POST"],
    allow_headers=["Authorization", "Content-Type"],
)


class UploadResponse(BaseModel):
    saved_path: str
    extracted_text: str
    char_count: int
    file_size: int
    mime: str
    sha256: str


@app.get("/healthz")
def healthz() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/upload", response_model=UploadResponse)
async def upload(
    file: Annotated[UploadFile, File()],
    category: Annotated[str, Form()] = DEFAULT_CATEGORY,
    _auth: None = Depends(require_auth),
) -> UploadResponse:
    if not file.filename:
        raise HTTPException(status_code=400, detail="缺少 filename")

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"目前只支援 {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # 讀整檔到記憶體（50MB 上限）
    data = await file.read()
    if len(data) == 0:
        raise HTTPException(status_code=400, detail="檔案是空的")
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"檔案 {len(data) / 1024 / 1024:.1f}MB 超過 {MAX_FILE_BYTES // 1024 // 1024}MB 上限",
        )

    # 抽文字（解析失敗的話檔案不寫，直接回錯誤）
    extracted = extract(data, ext)
    if not extracted.strip():
        log.warning("Extracted text empty for %s (likely scanned PDF)", file.filename)

    # 寫檔到 ruru 論文資料夾
    safe_category = sanitize_filename(category) or DEFAULT_CATEGORY
    target_dir = UPLOAD_ROOT / safe_category
    target_dir.mkdir(parents=True, exist_ok=True)

    safe_name = sanitize_filename(file.filename)
    target_path = reserve_path(target_dir, safe_name)
    target_path.write_bytes(data)

    sha256 = hashlib.sha256(data).hexdigest()
    log.info(
        "Uploaded %s (%d bytes, %d chars) → %s",
        safe_name, len(data), len(extracted), target_path,
    )

    return UploadResponse(
        saved_path=str(target_path),
        extracted_text=extracted,
        char_count=len(extracted),
        file_size=len(data),
        mime=file.content_type or "application/octet-stream",
        sha256=sha256,
    )


class FileEntry(BaseModel):
    name: str
    path: str
    category: str
    size: int
    modified: str


class ListResponse(BaseModel):
    root: str
    files: list[FileEntry]


@app.get("/list", response_model=ListResponse)
def list_files(_auth: None = Depends(require_auth)) -> ListResponse:
    if not UPLOAD_ROOT.exists():
        return ListResponse(root=str(UPLOAD_ROOT), files=[])
    entries: list[FileEntry] = []
    for path in UPLOAD_ROOT.rglob("*"):
        if not path.is_file():
            continue
        if path.name.startswith("."):
            continue
        rel = path.relative_to(UPLOAD_ROOT)
        category = rel.parts[0] if len(rel.parts) > 1 else ""
        stat = path.stat()
        entries.append(
            FileEntry(
                name=path.name,
                path=str(path),
                category=category,
                size=stat.st_size,
                modified=datetime.fromtimestamp(stat.st_mtime).isoformat(),
            )
        )
    entries.sort(key=lambda e: e.modified, reverse=True)
    return ListResponse(root=str(UPLOAD_ROOT), files=entries)


@app.get("/file/text")
def read_file_text(
    path: str,
    _auth: None = Depends(require_auth),
) -> dict[str, str | int]:
    """重讀已上傳檔案的文字（不重存 — 純抽文字回傳給前端）。
    path 必須是絕對路徑且在 UPLOAD_ROOT 之下，避免任意路徑讀取。"""
    target = Path(path).resolve()
    try:
        target.relative_to(UPLOAD_ROOT.resolve())
    except ValueError as e:
        raise HTTPException(status_code=403, detail="path 不在論文資料夾內") from e
    if not target.is_file():
        raise HTTPException(status_code=404, detail="檔案不存在")
    ext = target.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支援 {ext}")
    data = target.read_bytes()
    text = extract(data, ext)
    return {"path": str(target), "extracted_text": text, "char_count": len(text)}


if __name__ == "__main__":
    import uvicorn

    log.info("Token loaded; share via tunnel + frontend env")
    log.info("UPLOAD_ROOT = %s", UPLOAD_ROOT)
    uvicorn.run(app, host="127.0.0.1", port=9120, log_level="info")
